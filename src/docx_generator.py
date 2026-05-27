import asyncio
import uuid
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt

from utils.loggers import logger


SUPPORTED_MARKDOWN_EXTENSIONS = ["extra", "sane_lists"]


def _add_inline_runs(paragraph, node):
    for child in node.children:
        if isinstance(child, NavigableString):
            paragraph.add_run(str(child))
            continue

        if not isinstance(child, Tag):
            continue

        text = child.get_text()
        if child.name == "strong":
            run = paragraph.add_run(text)
            run.bold = True
        elif child.name == "em":
            run = paragraph.add_run(text)
            run.italic = True
        elif child.name == "code":
            run = paragraph.add_run(text)
            run.font.name = "Courier New"
        elif child.name == "a":
            run = paragraph.add_run(text)
            run.underline = True
            if child.get("href"):
                paragraph.add_run(f" ({child['href']})")
        else:
            _add_inline_runs(paragraph, child)


def _add_paragraph(document: Document, node: Tag):
    paragraph = document.add_paragraph()
    _add_inline_runs(paragraph, node)


def _add_list(document: Document, node: Tag, level: int = 0):
    style = "List Number" if node.name == "ol" else "List Bullet"

    for li in node.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        if level:
            paragraph.paragraph_format.left_indent = Pt(18 * level)

        for child in li.children:
            if isinstance(child, NavigableString):
                paragraph.add_run(str(child))
            elif isinstance(child, Tag) and child.name in {"ul", "ol"}:
                _add_list(document, child, level + 1)
            elif isinstance(child, Tag):
                if child.name == "p":
                    _add_inline_runs(paragraph, child)
                else:
                    _add_inline_runs(paragraph, child)


def _add_node(document: Document, node):
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            document.add_paragraph(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "h1":
        document.add_heading(node.get_text(" ", strip=True), level=1)
    elif node.name == "h2":
        document.add_heading(node.get_text(" ", strip=True), level=2)
    elif node.name == "h3":
        document.add_heading(node.get_text(" ", strip=True), level=3)
    elif node.name == "p":
        _add_paragraph(document, node)
    elif node.name in {"ul", "ol"}:
        _add_list(document, node)
    elif node.name == "hr":
        document.add_paragraph("_" * 72)
    elif node.name == "blockquote":
        paragraph = document.add_paragraph(style="Intense Quote")
        _add_inline_runs(paragraph, node)
    else:
        for child in node.children:
            _add_node(document, child)


def _build_docx(report: str, file_path: Path):
    html_report = markdown.markdown(report, extensions=SUPPORTED_MARKDOWN_EXTENSIONS)
    soup = BeautifulSoup(html_report, "html.parser")
    document = Document()

    section = document.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    for child in soup.children:
        _add_node(document, child)

    document.save(file_path)


async def generate_docx(report: str) -> str:
    logger.info("Starting DOCX generation...")

    if not report or not report.strip():
        raise ValueError("Report content is empty. Cannot generate DOCX.")

    try:
        output_dir = Path("docx")
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / f"report_{uuid.uuid4().hex}.docx"
        await asyncio.to_thread(_build_docx, report, file_path)
        logger.info("DOCX generation completed successfully. File saved as %s", file_path)
        return str(file_path)
    except Exception as e:
        logger.error("Failed to generate DOCX: %s", e)
        raise RuntimeError(f"Failed to generate DOCX: {e}")
